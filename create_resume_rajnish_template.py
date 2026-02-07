#!/usr/bin/env python3
"""
Script to create a two-column DOCX resume matching the Rajnish-style template:
centered header (name, title, contact with icons) and two white columns with
serif body text and sans-serif section headers. Design has no outer page border;
if you see a frame when exporting, disable border/frame in the export dialog.
Optional: for a very light grey/off-white background, set page color in Word
after opening the DOCX (Design > Page Color).

Output: Rajnish-Kumar.docx
To get Rajnish-Kumar.pdf, open Rajnish-Kumar.docx and export to PDF, or run
convert_to_pdf.py after editing it to use input Rajnish-Kumar.docx and output
Rajnish-Kumar.pdf.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
import os

# Script-relative base path for icons
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Fonts: Use Times New Roman throughout (serif font matching the design)
SERIF_FONT = 'Times New Roman'
SANS_SERIF_FONT = 'Times New Roman'  # Use same font throughout

# Dark teal-green for section header icons and borders (matches reference image)
SECTION_ACCENT_COLOR = '2D5A5A'  # OOXML hex (no #)
SECTION_ACCENT_COLOR_URL = '%232D5A5A'  # URL-encoded for Iconify

# Icon file mappings - PNG files in icon folder
ICON_FILES = {
    'profile': 'profile.png',
    'skills': 'skill.png',
    'education': 'education.png',
    'certificates': 'certificate.png',
    'findmeonline': 'findme.png',
    'experience': 'experience.png',
    'projects': 'project.png',
    'external_link': 'link.png',  # Using link.png as external link icon
    'location': 'location-pin.png',
    'email': 'email.png',
    'phone': 'call.png',
}

def get_icon_path(icon_key):
    """Get full path to icon PNG file."""
    if icon_key not in ICON_FILES:
        return None
    icon_file = ICON_FILES[icon_key]
    icon_path = os.path.join(SCRIPT_DIR, 'icon', icon_file)
    if os.path.exists(icon_path):
        return icon_path
    return None

def set_margins(section, top, bottom, left, right):
    """Set page margins for a section."""
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def _add_short_bold_line(cell):
    """Add a short, bold horizontal line below section header."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)  # No gap - line directly below title
    # p.paragraph_format.space_after = Pt(4)
    # p.paragraph_format.line_spacing = 1.0  # Tight line spacing
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Inches(3)  # Make it shorter (not full width)
    
    # Add bold bottom border
    pPr = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Thicker line (12 = 1.5pt, was 6 = 0.75pt)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SECTION_ACCENT_COLOR)  # Dark teal-green to match icons
    pBdr.append(bottom)
    pPr.append(pBdr)
    p._element.insert(0, pPr)
    return p

def add_section_header(cell, text, icon_key=None, is_left_column=False):
    """Add a section header: optional icon, bold uppercase sans-serif title, then short bold line."""
    # Text paragraph - no right indent to prevent vertical wrapping
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)  # No gap before border line
    p.paragraph_format.line_spacing = 1.0  # Tight line spacing
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    # NO right_indent here - it causes text to wrap vertically
    
    if icon_key:
        icon_path = get_icon_path(icon_key)
        if icon_path:
            try:
                # Ensure no left indent for both columns
                pPr = p._element.get_or_add_pPr()
                # Remove existing indentation if present
                existing_ind = pPr.find(qn('w:ind'))
                if existing_ind is not None:
                    pPr.remove(existing_ind)
                
                # Only apply negative indent for left column to offset its cell margin
                # Right column has 0 left margin, so no negative indent needed
                if is_left_column:
                    # Create indentation element with smaller negative indent to move icons more to the right
                    # -144 twips = -0.1 inches (smaller offset to move icons more to the right)
                    # This keeps icons away from the left edge while maintaining proper alignment
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), '-144')  # Smaller negative indent to move icons more to the right
                    ind.set(qn('w:firstLine'), '0')
                    pPr.append(ind)
                else:
                    # For right column, explicitly set left indent to 0 to remove any default spacing
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), '-144')
                    ind.set(qn('w:firstLine'), '0')
                    pPr.append(ind)
                
                icon_run = p.add_run()
                icon_run.add_picture(icon_path, width=Pt(10), height=Pt(10))
                # Minimal spacing between icon and text - using very small non-breaking space
                sp = p.add_run('\u2009')  # Thin space character (smaller than regular space)
                sp.font.size = Pt(0.5)  # Very small spacing to minimize gap
            except Exception as e:
                print(f"Warning: Could not add icon {icon_key}: {e}")
                pass
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = SANS_SERIF_FONT  # Use same font as all other content
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Border line paragraph - separate paragraph with right indent for shorter line
    border_p = cell.add_paragraph()
    border_p.paragraph_format.space_before = Pt(0)  # No gap - directly below title
    border_p.paragraph_format.space_after = Pt(4)  # Space after the border line
    border_p.paragraph_format.line_spacing = 1.0
    border_p.paragraph_format.left_indent = Pt(0)
    # Adjust right indent based on column - right column should match icon+title length
    if is_left_column:
        border_p.paragraph_format.right_indent = Inches(1.5)  # Left column: shorter border
    else:
        border_p.paragraph_format.right_indent = Inches(0.5)  # Right column: longer border to match icon+title
    
    # Add a non-breaking space to give the paragraph width so the border renders properly
    border_run = border_p.add_run('\u00A0')  # Non-breaking space
    border_run.font.size = Pt(1)  # Very small, invisible content
    
    # Add bold bottom border to this paragraph
    pPr = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Thicker line (12 = 1.5pt)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SECTION_ACCENT_COLOR)  # Dark teal-green to match icons
    pBdr.append(bottom)
    pPr.append(pBdr)
    border_p._element.insert(0, pPr)
    
    return p

def add_body_text(cell, text, size=Pt(9)):
    """Add body paragraph, black text (serif)."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = SANS_SERIF_FONT
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    return p

def add_bullet(cell, text, size=Pt(9)):
    """Add bullet point with optional **bold** keywords (bold in black)."""
    p = cell.add_paragraph()
    bullet_run = p.add_run('• ')
    bullet_run.font.size = size
    bullet_run.font.name = SANS_SERIF_FONT
    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
    bullet_run.bold = True
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.line_spacing = 1.05
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.size = size
                run.font.name = SANS_SERIF_FONT
                run.font.color.rgb = RGBColor(0, 0, 0)
                if i % 2 == 1:
                    run.bold = True
    else:
        run = p.add_run(text)
        run.font.size = size
        run.font.name = SANS_SERIF_FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_contact_header_centered(cell_or_doc, icon_dir, location, email, phone, icon_size=Pt(9)):
    """Add centered contact line with inline icons (location, email, phone)."""
    p = cell_or_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)  # Increased from Pt(4) for more height
    items = [
        ('location', location),
        ('email', email),
        ('phone', phone),
    ]
    for i, (icon_key, text) in enumerate(items):
        icon_path = get_icon_path(icon_key)
        if icon_path:
            try:
                icon_run = p.add_run()
                icon_run.add_picture(icon_path, width=icon_size, height=icon_size)
                # Reduced spacing between icon and text
                sp = p.add_run('\u2009')  # Thin space character
                sp.font.size = Pt(2)  # Reduced from Pt(6) to minimize gap
            except Exception as e:
                print(f"Warning: Could not add icon {icon_key}: {e}")
                pass
        txt = p.add_run(text)
        txt.font.size = Pt(10)
        txt.font.name = SANS_SERIF_FONT
        txt.font.color.rgb = RGBColor(0, 0, 0)
        if '@' in text:
            txt.element.set(qn('w:noBreak'), '1')
        if i < len(items) - 1:
            sep = p.add_run('  ')  # Reduced spacing between contact items from '   ' to '  '
            sep.font.size = Pt(4)  # Reduced from Pt(6)
    return p

def set_cell_margins(cell, top, bottom, left, right):
    """Set cell margins (inches)."""
    tc = cell._element
    tcPr = tc.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.append(tcPr)
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(int(margin_value * 1440)))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def create_rajnish_resume():
    """Build the Rajnish-style two-column resume (no page border; light grey page background)."""
    doc = Document()
    section = doc.sections[0]
    set_margins(section, 0, 0, 0, 0)  # No page margins - content fits end-to-end
    
    # Set page background to very light grey
    sectPr = section._sectPr
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), 'F8F8F8')  # Very light grey
    sectPr.append(bg)

    style = doc.styles['Normal']
    style.font.name = SANS_SERIF_FONT
    style.font.size = Pt(10.5)

    # ----- Header: name, title, contact (with background color) -----
    # Create header table with background color
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Set table width to full page width (8.5" page)
    header_table.columns[0].width = Inches(8.5)

    # Set background color for header cell
    header_tcPr = header_cell._element.tcPr
    if header_tcPr is None:
        header_tcPr = OxmlElement('w:tcPr')
        header_cell._element.insert(0, header_tcPr)
    header_shading = OxmlElement('w:shd')
    header_shading.set(qn('w:fill'), 'F5F5F5')  # Very light grey
    header_tcPr.append(header_shading)

    # Remove table borders
    header_tbl = header_table._element
    header_tblPr = header_tbl.tblPr
    if header_tblPr is None:
        header_tblPr = OxmlElement('w:tblPr')
        header_tbl.insert(0, header_tblPr)
    header_tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        header_tblBorders.append(border)
    header_tblPr.append(header_tblBorders)

    # Set cell margins - increased top and bottom for more height
    set_cell_margins(header_cell, 0.30, 0.40, 0.0, 0.0)  # Increased top to 0.15", bottom to 0.25"

    # Clear default paragraph and add header content
    header_cell.paragraphs[0].clear()

    name_p = header_cell.add_paragraph()
    name_p.paragraph_format.space_before = Pt(4)  # Add some top spacing
    name_run = name_p.add_run('Rajnish Kumar')
    name_run.bold = True
    name_run.font.size = Pt(19)
    name_run.font.name = SANS_SERIF_FONT
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(4)  # Increased from Pt(2)

    title_p = header_cell.add_paragraph()
    title_run = title_p.add_run('Senior Software Engineer')
    title_run.font.size = Pt(12)
    title_run.font.name = SANS_SERIF_FONT
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)  # Increased from Pt(4)

    add_contact_header_centered(
        header_cell, SCRIPT_DIR,
        'Gurugram, IN',
        'kumar89rajnish@gmail.com',
        '+91 8506981317',
    )
    
    # Add spacing after header (margin at bottom of header)
    if header_cell.paragraphs:
        header_cell.paragraphs[-1].paragraph_format.space_after = Pt(10)  # Increased from Pt(8) for more height

    # ----- Two-column table -----
    table = doc.add_table(rows=1, cols=2)
    left_col = table.columns[0]
    right_col = table.columns[1]
    # Full page width (8.5") - equal 50/50 split
    left_col.width = Inches(3.90)
    right_col.width = Inches(4.60)

    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    # Add left padding to prevent text cutoff, minimal other margins
    set_cell_margins(left_cell, 0.0, 0.0, 0.55, 0.2)  # Increased left padding
    set_cell_margins(right_cell, 0.0, 0.0, 0.0, 0.2)  # Small left padding, no right margin
    
    # Set white background for both columns to contrast with page background
    left_tcPr = left_cell._element.tcPr
    if left_tcPr is None:
        left_tcPr = OxmlElement('w:tcPr')
        left_cell._element.insert(0, left_tcPr)
    shading_left = OxmlElement('w:shd')
    shading_left.set(qn('w:fill'), 'FFFFFF')  # White
    left_tcPr.append(shading_left)
    
    right_tcPr = right_cell._element.tcPr
    if right_tcPr is None:
        right_tcPr = OxmlElement('w:tcPr')
        right_cell._element.insert(0, right_tcPr)
    shading_right = OxmlElement('w:shd')
    shading_right.set(qn('w:fill'), 'FFFFFF')  # White
    right_tcPr.append(shading_right)

    tr = table.rows[0]._element
    trPr = tr.trPr
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '0')
    trHeight.set(qn('w:hRule'), 'auto')
    trPr.append(trHeight)

    # Clear default paragraph in cells
    left_cell.paragraphs[0].clear()
    right_cell.paragraphs[0].clear()

    # ========== LEFT COLUMN ==========
    add_section_header(left_cell, 'Profile', icon_key='profile', is_left_column=True)
    add_body_text(
        left_cell,
        'Software Engineer with 6 years of experience in mobile app development, specializing in React Native, JavaScript, and TypeScript. Expertise in full-stack development, native module integration, and leading cross-functional teams to deliver scalable solutions in agile environments.'
    )

    add_section_header(left_cell, 'Skills', icon_key='skills', is_left_column=True)
    # App Development
    sub = left_cell.add_paragraph()
    sub_run = sub.add_run('App Development')
    sub_run.bold = True
    sub_run.font.size = Pt(10)
    sub_run.font.name = SANS_SERIF_FONT
    sub_run.font.color.rgb = RGBColor(0, 0, 0)
    sub.paragraph_format.space_after = Pt(1)
    tech_p = left_cell.add_paragraph()
    tech_prefix = tech_p.add_run('Tech:')
    tech_prefix.bold = True
    tech_prefix.font.size = Pt(8.5)
    tech_prefix.font.name = SANS_SERIF_FONT
    tech_prefix.font.color.rgb = RGBColor(0, 0, 0)
    tech_content = tech_p.add_run(' React Native, React.js, Redux, Kotlin, Swift, Reanimated Animations, Native Navigation, Gradle, Xcode, Custom Native Module Development')
    tech_content.font.size = Pt(8.5)
    tech_content.font.name = SANS_SERIF_FONT
    tech_content.font.color.rgb = RGBColor(0, 0, 0)
    tech_p.paragraph_format.space_after = Pt(2)
    
    modules_p = left_cell.add_paragraph()
    modules_prefix = modules_p.add_run('Modules:')
    modules_prefix.bold = True
    modules_prefix.font.size = Pt(8.5)
    modules_prefix.font.name = SANS_SERIF_FONT
    modules_prefix.font.color.rgb = RGBColor(0, 0, 0)
    modules_content = modules_p.add_run(' Payment Integration, Subscription Management')
    modules_content.font.size = Pt(8.5)
    modules_content.font.name = SANS_SERIF_FONT
    modules_content.font.color.rgb = RGBColor(0, 0, 0)
    modules_p.paragraph_format.space_after = Pt(2)

    sub2 = left_cell.add_paragraph()
    sub2_run = sub2.add_run('Fullstack Development')
    sub2_run.bold = True
    sub2_run.font.size = Pt(9)
    sub2_run.font.name = SANS_SERIF_FONT
    sub2_run.font.color.rgb = RGBColor(0, 0, 0)
    sub2.paragraph_format.space_before = Pt(4)
    sub2.paragraph_format.space_after = Pt(1)
    frontend_p = left_cell.add_paragraph()
    frontend_prefix = frontend_p.add_run('Frontend:')
    frontend_prefix.bold = True
    frontend_prefix.font.size = Pt(8.5)
    frontend_prefix.font.name = SANS_SERIF_FONT
    frontend_prefix.font.color.rgb = RGBColor(0, 0, 0)
    frontend_content = frontend_p.add_run(' React, Redux, HTML, CSS, JavaScript')
    frontend_content.font.size = Pt(8.5)
    frontend_content.font.name = SANS_SERIF_FONT
    frontend_content.font.color.rgb = RGBColor(0, 0, 0)
    frontend_p.paragraph_format.space_after = Pt(2)
    
    backend_p = left_cell.add_paragraph()
    backend_prefix = backend_p.add_run('Backend:')
    backend_prefix.bold = True
    backend_prefix.font.size = Pt(8.5)
    backend_prefix.font.name = SANS_SERIF_FONT
    backend_prefix.font.color.rgb = RGBColor(0, 0, 0)
    backend_content = backend_p.add_run(' Node.js, Express, TypeScript, MongoDB')
    backend_content.font.size = Pt(8.5)
    backend_content.font.name = SANS_SERIF_FONT
    backend_content.font.color.rgb = RGBColor(0, 0, 0)
    backend_p.paragraph_format.space_after = Pt(2)

    sub3 = left_cell.add_paragraph()
    sub3_run = sub3.add_run('Tools & Technologies')
    sub3_run.bold = True
    sub3_run.font.size = Pt(9)
    sub3_run.font.name = SANS_SERIF_FONT
    sub3_run.font.color.rgb = RGBColor(0, 0, 0)
    sub3.paragraph_format.space_before = Pt(4)
    sub3.paragraph_format.space_after = Pt(1)
    devtools_p = left_cell.add_paragraph()
    devtools_prefix = devtools_p.add_run('Development Tools:')
    devtools_prefix.bold = True
    devtools_prefix.font.size = Pt(8.5)
    devtools_prefix.font.name = SANS_SERIF_FONT
    devtools_prefix.font.color.rgb = RGBColor(0, 0, 0)
    devtools_content = devtools_p.add_run(' Android Studio, Xcode, Flipper, Strapi (CMS)')
    devtools_content.font.size = Pt(8.5)
    devtools_content.font.name = SANS_SERIF_FONT
    devtools_content.font.color.rgb = RGBColor(0, 0, 0)
    devtools_p.paragraph_format.space_after = Pt(2)
    
    collab_p = left_cell.add_paragraph()
    collab_prefix = collab_p.add_run('Collaboration & Design:')
    collab_prefix.bold = True
    collab_prefix.font.size = Pt(8.5)
    collab_prefix.font.name = SANS_SERIF_FONT
    collab_prefix.font.color.rgb = RGBColor(0, 0, 0)
    collab_content = collab_p.add_run(' Git, Postman, Firebase, Figma, JIRA')
    collab_content.font.size = Pt(8.5)
    collab_content.font.name = SANS_SERIF_FONT
    collab_content.font.color.rgb = RGBColor(0, 0, 0)
    collab_p.paragraph_format.space_after = Pt(2)

    sub4 = left_cell.add_paragraph()
    sub4_run = sub4.add_run('Others')
    sub4_run.bold = True
    sub4_run.font.size = Pt(9)
    sub4_run.font.name = SANS_SERIF_FONT
    sub4_run.font.color.rgb = RGBColor(0, 0, 0)
    sub4.paragraph_format.space_before = Pt(4)
    sub4.paragraph_format.space_after = Pt(1)
    api_p = left_cell.add_paragraph()
    api_prefix = api_p.add_run('API & Server:')
    api_prefix.bold = True
    api_prefix.font.size = Pt(8.5)
    api_prefix.font.name = SANS_SERIF_FONT
    api_prefix.font.color.rgb = RGBColor(0, 0, 0)
    api_content = api_p.add_run(' Swagger, REST API, NGINX, S3 Browser')
    api_content.font.size = Pt(8.5)
    api_content.font.name = SANS_SERIF_FONT
    api_content.font.color.rgb = RGBColor(0, 0, 0)
    api_p.paragraph_format.space_after = Pt(2)

    add_section_header(left_cell, 'Education', icon_key='education', is_left_column=True)
    # B.Tech
    btech_label = left_cell.add_paragraph()
    btech_label_run = btech_label.add_run('B.Tech')
    btech_label_run.bold = True
    btech_label_run.font.size = Pt(8.5)
    btech_label_run.font.name = SANS_SERIF_FONT
    btech_label_run.font.color.rgb = RGBColor(0, 0, 0)
    btech_label.paragraph_format.space_after = Pt(1)
    add_body_text(left_cell, 'KIIT College of Engineering', Pt(8.5))
    add_body_text(left_cell, '2016 – 2020 | Sohna, Gurugram, India', Pt(8.5))
    add_body_text(left_cell, 'Computer science and Engineering', Pt(8.5))
    
    # Diploma
    diploma_label = left_cell.add_paragraph()
    diploma_label.paragraph_format.space_before = Pt(4)
    diploma_label_run = diploma_label.add_run('Diploma')
    diploma_label_run.bold = True
    diploma_label_run.font.size = Pt(8.5)
    diploma_label_run.font.name = SANS_SERIF_FONT
    diploma_label_run.font.color.rgb = RGBColor(0, 0, 0)
    diploma_label.paragraph_format.space_after = Pt(1)
    add_body_text(left_cell, 'Future Soft Technology', Pt(8.5))
    add_body_text(left_cell, "2015 – 2016 | Gurugram, India", Pt(8.5))
    add_body_text(left_cell, "DOEACC 'O' Level Course", Pt(8.5))

    add_section_header(left_cell, 'Certificates', icon_key='certificates', is_left_column=True)
    udemy_p = left_cell.add_paragraph()
    # Add hyperlink to Udemy text
    udemy_url = 'https://www.udemy.com/certificate/UC-4994896d-7582-4baa-9b1a-691fec35c68d/'
    # Create hyperlink using OOXML
    rId = udemy_p.part.relate_to(udemy_url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rId)
    # Create run inside hyperlink
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Set font properties
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')  # Black color
    rPr.append(color)
    b = OxmlElement('w:b')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '17')  # 8.5pt = 17 half-points
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), SANS_SERIF_FONT)
    rPr.append(rFonts)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = 'Udemy'
    run.append(t)
    hyperlink.append(run)
    udemy_p._element.append(hyperlink)
    # Add external link icon with hyperlink
    external_link_path = get_icon_path('external_link')
    if external_link_path:
        try:
            udemy_p.add_run(' ')
            # Create hyperlink for the icon
            icon_rId = udemy_p.part.relate_to(udemy_url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            icon_hyperlink = OxmlElement('w:hyperlink')
            icon_hyperlink.set(qn('r:id'), icon_rId)
            # Create run for the image inside hyperlink
            icon_run_elem = OxmlElement('w:r')
            icon_hyperlink.append(icon_run_elem)
            # Add picture to a temporary run first to get the drawing element
            temp_run = udemy_p.add_run()
            temp_run.add_picture(external_link_path, width=Pt(7), height=Pt(7))
            # Move the drawing element from temp_run to icon_run inside hyperlink
            drawing_elem = temp_run._element.find(qn('w:drawing'))
            if drawing_elem is not None:
                temp_run._element.remove(drawing_elem)
                icon_run_elem.append(drawing_elem)
            # Remove the temporary run if it's empty
            if len(temp_run._element) == 0:
                udemy_p._element.remove(temp_run._element)
            # Append hyperlink to paragraph
            udemy_p._element.append(icon_hyperlink)
        except Exception as e:
            print(f"Warning: Could not add external link icon: {e}")
            pass
    udemy_p.paragraph_format.space_after = Pt(2)

    add_section_header(left_cell, 'Find Me Online', icon_key='findmeonline', is_left_column=True)
    add_body_text(left_cell, 'github.com/rajnish05', Pt(9.5))
    add_body_text(left_cell, 'linkedin.com/in/rajnish05', Pt(9.5))
    add_body_text(left_cell, 'stackoverflow.com/users/13081559/kumar', Pt(9.5))
    add_body_text(left_cell, 'npmjs.com/kumar_01', Pt(9.5))

    # ========== RIGHT COLUMN ==========
    add_section_header(right_cell, 'Professional Experience', icon_key='experience', is_left_column=False)

    # Zupee - App Developer (Most recent)
    exp1_company = right_cell.add_paragraph()
    exp1_run = exp1_company.add_run('Zupee')
    exp1_run.bold = True
    exp1_run.font.size = Pt(11)
    exp1_run.font.name = SANS_SERIF_FONT
    exp1_run.font.color.rgb = RGBColor(0, 0, 0)
    exp1_company.paragraph_format.space_after = Pt(0)
    
    exp1_title = right_cell.add_paragraph()
    exp1_title_run = exp1_title.add_run('App Developer')
    exp1_title_run.bold = True
    exp1_title_run.font.size = Pt(10)
    exp1_title_run.font.name = SANS_SERIF_FONT
    exp1_title_run.font.color.rgb = RGBColor(0, 0, 0)
    exp1_title.paragraph_format.space_after = Pt(0)
    
    exp1_line = right_cell.add_paragraph()
    exp1_line_run = exp1_line.add_run('Apr 2025 – present | Gurugram, India')
    exp1_line_run.font.size = Pt(10)
    exp1_line_run.font.name = SANS_SERIF_FONT
    exp1_line_run.font.color.rgb = RGBColor(80, 80, 80)
    exp1_line.paragraph_format.space_after = Pt(2)
    add_bullet(right_cell, 'Developed real-time voice call module using **WebRTC** and native audio APIs, ensuring low-latency communication for gaming platform')
    add_bullet(right_cell, 'Implemented voice notes feature with audio recording, compression, and playback functionality, improving user engagement by 40%')
    add_bullet(right_cell, 'Built scalable chat system with message queuing, real-time synchronization, and multimedia support (text, images, voice notes)')
    add_bullet(right_cell, 'Optimized communication modules for performance, reducing memory footprint by 30% and improving app responsiveness')

    # Dresma AI - SDE2 (Consolidated)
    exp2_company = right_cell.add_paragraph()
    exp2_run = exp2_company.add_run('Dresma AI')
    exp2_run.bold = True
    exp2_run.font.size = Pt(11)
    exp2_run.font.name = SANS_SERIF_FONT
    exp2_run.font.color.rgb = RGBColor(0, 0, 0)
    exp2_company.paragraph_format.space_before = Pt(4)
    exp2_company.paragraph_format.space_after = Pt(0)
    
    exp2_title = right_cell.add_paragraph()
    exp2_title_run = exp2_title.add_run('SDE2')
    exp2_title_run.bold = True
    exp2_title_run.font.size = Pt(10)
    exp2_title_run.font.name = SANS_SERIF_FONT
    exp2_title_run.font.color.rgb = RGBColor(0, 0, 0)
    exp2_title.paragraph_format.space_after = Pt(0)
    
    exp2_line = right_cell.add_paragraph()
    exp2_line_run = exp2_line.add_run('Apr 2020 – Apr 2025 | Gurugram, India')
    exp2_line_run.font.size = Pt(10)
    exp2_line_run.font.name = SANS_SERIF_FONT
    exp2_line_run.font.color.rgb = RGBColor(80, 80, 80)
    exp2_line.paragraph_format.space_after = Pt(2)
    add_bullet(right_cell, 'Led **React Native** development initiatives, managing feature lifecycle from design to deployment across iOS and Android platforms')
    add_bullet(right_cell, 'Architected and developed mobile application from ground up, resulting in 30% revenue increase and seamless App Store/Play Store deployments')
    add_bullet(right_cell, 'Implemented payment gateway integrations (Paddle, Stripe) and subscription management systems, enhancing in-app purchase capabilities')
    add_bullet(right_cell, 'Collaborated with product and engineering teams on API design, automated testing frameworks, and code quality standards, reducing bugs by 25%')

    # Differential Edge
    exp3_company = right_cell.add_paragraph()
    exp3_run = exp3_company.add_run('Differential Edge Pvt. Ltd')
    exp3_run.bold = True
    exp3_run.font.size = Pt(11)
    exp3_run.font.name = SANS_SERIF_FONT
    exp3_run.font.color.rgb = RGBColor(0, 0, 0)
    exp3_company.paragraph_format.space_before = Pt(4)
    exp3_company.paragraph_format.space_after = Pt(0)
    
    exp3_title = right_cell.add_paragraph()
    exp3_title_run = exp3_title.add_run('MERN stack Developer')
    exp3_title_run.bold = True
    exp3_title_run.font.size = Pt(10)
    exp3_title_run.font.name = SANS_SERIF_FONT
    exp3_title_run.font.color.rgb = RGBColor(0, 0, 0)
    exp3_title.paragraph_format.space_after = Pt(0)
    
    exp3_line = right_cell.add_paragraph()
    exp3_line_run = exp3_line.add_run('Sep 2019 – Mar 2020 | Gurugram, India')
    exp3_line_run.font.size = Pt(10)
    exp3_line_run.font.name = SANS_SERIF_FONT
    exp3_line_run.font.color.rgb = RGBColor(80, 80, 80)
    exp3_line.paragraph_format.space_after = Pt(2)
    add_bullet(right_cell, 'Integrated **authentication REST APIs** using Node, Express, and React for a seamless user experience.')
    add_bullet(right_cell, 'Documented APIs using **Swagger** with Node, Express, and JSON, delivering clear and comprehensive documentation for efficient development.')
    add_bullet(right_cell, 'Developed a core component for an **in-house project**, laying a solid foundation for future feature implementation and expansion.')

    add_section_header(right_cell, 'Projects', icon_key='projects', is_left_column=False)

    proj_title = right_cell.add_paragraph()
    proj_run = proj_title.add_run('FindMyHome')
    proj_run.bold = True
    proj_run.font.size = Pt(11)
    proj_run.font.name = SANS_SERIF_FONT
    proj_run.font.color.rgb = RGBColor(0, 0, 0)
    # Add external link icon
    external_link_path = get_icon_path('external_link')
    proj_title.paragraph_format.space_after = Pt(0)
    proj_tech = right_cell.add_paragraph()
    tech_label_run = proj_tech.add_run('Tech Stack:')
    tech_label_run.bold = True
    tech_label_run.font.size = Pt(9.5)
    tech_label_run.font.name = SANS_SERIF_FONT
    tech_label_run.font.color.rgb = RGBColor(0, 0, 0)
    tech_content_run = proj_tech.add_run(' React-Native, TypeScript, React, Redux, Firebase,')
    tech_content_run.font.size = Pt(9.5)
    tech_content_run.font.name = SANS_SERIF_FONT
    tech_content_run.font.color.rgb = RGBColor(0, 0, 0)
    proj_tech.paragraph_format.space_after = Pt(2)
    
    key_features_label = right_cell.add_paragraph()
    kf_run = key_features_label.add_run('Key Features:')
    kf_run.bold = True
    kf_run.font.size = Pt(10)
    kf_run.font.name = SANS_SERIF_FONT
    kf_run.font.color.rgb = RGBColor(0, 0, 0)
    key_features_label.paragraph_format.space_after = Pt(1)
    
    add_bullet(right_cell, '**House Listings and Details**: The app allows users to browse a list of houses and view detailed information about each property.')
    add_bullet(right_cell, '**Authentication Flow**: It includes a login flow for user authentication.')
    add_bullet(right_cell, "**Location-Based Features**: The app unlocks detailed house information based on the user's proximity to specific locations.")

    # Save
    output_file = os.path.join(SCRIPT_DIR, 'Rajnish-Kumar.docx')
    doc.save(output_file)
    print(f"Resume created: {output_file}")
    print("To get Rajnish-Kumar.pdf, open the DOCX in Word and export to PDF, or run convert_to_pdf.py with this file.")


if __name__ == '__main__':
    try:
        create_rajnish_resume()
    except ImportError:
        print("Error: python-docx not found. Install with: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
