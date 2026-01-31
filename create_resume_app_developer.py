#!/usr/bin/env python3
"""
Script to create a two-column single-page ATS-friendly DOCX resume
with modern sidebar design.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_margins(section, top, bottom, left, right):
    """Set page margins for a section"""
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_sidebar_header(cell, text, is_main=False):
    """Add a sidebar section header to a cell"""
    p = cell.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    if is_main:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_main_header(cell, text):
    """Add a main section header to a cell"""
    p = cell.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    return p

def add_sidebar_text(cell, text, size=Pt(9), color=RGBColor(255, 255, 255)):
    """Add text to sidebar cell"""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    return p

def add_sidebar_text_with_icon(cell, icon_path, text, size=Pt(8.5), color=RGBColor(255, 255, 255), icon_size=Pt(6.5)):
    """Add text to sidebar cell with PNG icon"""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Prevent text wrapping - keep icon and text together
    p.paragraph_format.widow_control = False
    p.paragraph_format.keep_together = True
    
    # Add icon if file exists
    if os.path.exists(icon_path):
        try:
            icon_run = p.add_run()
            # Use smaller icon size for tighter spacing
            icon_run.add_picture(icon_path, width=icon_size, height=icon_size)
        except:
            # If image fails to load, just continue without icon
            pass
    
    # Add small space run with minimal width
    space_run = p.add_run(' ')
    space_run.font.size = Pt(3)
    
    # Add text - use non-breaking spaces for email to prevent wrapping
    if '@' in text:  # It's an email
        # Replace spaces with non-breaking spaces in email
        text = text.replace(' ', '\u00A0')
        text_run = p.add_run(text)
        text_run.font.size = size
        text_run.font.name = 'Calibri'
        text_run.font.color.rgb = color
        # Prevent breaking
        text_run.element.set(qn('w:noBreak'), '1')
    else:
        text_run = p.add_run(text)
        text_run.font.size = size
        text_run.font.name = 'Calibri'
        text_run.font.color.rgb = color
    
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Pt(0)
    return p

def add_sidebar_bullet(cell, text, size=Pt(9)):
    """Add bullet point to sidebar cell"""
    p = cell.add_paragraph()
    # Add bullet character manually
    bullet_run = p.add_run('• ')
    bullet_run.font.size = size
    bullet_run.font.name = 'Calibri'
    bullet_run.font.color.rgb = RGBColor(255, 255, 255)  # White bullet
    bullet_run.bold = True
    
    text_run = p.add_run(text)
    text_run.font.size = size
    text_run.font.name = 'Calibri'
    text_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.line_spacing = 1.05
    return p

def add_compact_bullet(cell, text):
    """Add compact bullet point to a cell"""
    p = cell.add_paragraph()
    # Add bullet character manually
    bullet_run = p.add_run('• ')
    bullet_run.font.size = Pt(9)
    bullet_run.font.name = 'Calibri'
    bullet_run.font.color.rgb = RGBColor(0, 0, 0)  # Black bullet
    bullet_run.bold = True
    
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.05
    
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    run = p.add_run(part)
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            else:
                run = p.add_run(part)
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for tech terms
    else:
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
    
    return p

def create_twocolumn_resume():
    """Create a two-column resume with sidebar"""
    
    doc = Document()
    
    # Set margins - reduced top and right margins
    section = doc.sections[0]
    set_margins(section, 0.2, 0.3, 0.3, 0.15)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9.5)
    
    # Create a table for two-column layout
    table = doc.add_table(rows=1, cols=2)
    
    # Set column widths: Left (sidebar) = 32%, Right (main) = 68%
    left_col = table.columns[0]
    right_col = table.columns[1]
    left_col.width = Inches(2.2)
    right_col.width = Inches(5.3)
    
    # Remove borders from table for cleaner look
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove all borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Get cells
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    
    # Set row height to allow content to expand
    tr = table.rows[0]._element
    trPr = tr.trPr
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    
    # Allow row to expand
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '0')  # Auto height
    trHeight.set(qn('w:hRule'), 'auto')
    trPr.append(trHeight)
    
    # Set sidebar background color (dark blue)
    left_tcPr = left_cell._element.tcPr
    if left_tcPr is None:
        left_tcPr = OxmlElement('w:tcPr')
        left_cell._element.insert(0, left_tcPr)
    
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '003366')  # Dark blue background
    left_tcPr.append(shading)
    
    # Set cell properties - ensure top alignment
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    
    # Set cell margins/padding
    def set_cell_margins(cell, top, bottom, left, right):
        """Set cell margins"""
        tc = cell._element
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), str(int(margin_value * 1440)))  # Convert inches to twips
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
    
    # Set padding for both cells (in inches) - reduced top margin further
    set_cell_margins(left_cell, 0.0, 0.1, 0.15, 0.1)
    set_cell_margins(right_cell, 0.0, 0.1, 0.15, 0.1)
    
    # ========== LEFT COLUMN (SIDEBAR) ==========
    # Clear the default paragraph
    left_para = left_cell.paragraphs[0]
    left_para.clear()
    
    # Name in sidebar
    name_para = left_cell.add_paragraph()
    name_run = name_para.add_run('PRABHAT KUMAR')
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.name = 'Calibri'
    name_run.font.color.rgb = RGBColor(255, 255, 255)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)
    
    # Title
    title_para = left_cell.add_paragraph()
    title_run = title_para.add_run('Application Developer\n(Mobile & Desktop)')
    title_run.font.size = Pt(10)
    title_run.font.name = 'Calibri'
    title_run.font.color.rgb = RGBColor(200, 220, 255)  # Light blue
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(5)
    
    # Divider line
    divider = left_cell.add_paragraph()
    divider_run = divider.add_run('─' * 25)
    divider_run.font.size = Pt(6)
    divider_run.font.color.rgb = RGBColor(200, 220, 255)
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.paragraph_format.space_after = Pt(4)
    
    # Contact Information - custom header with reduced spacing
    contact_header = left_cell.add_paragraph()
    contact_header_run = contact_header.add_run('CONTACT')
    contact_header_run.bold = True
    contact_header_run.font.size = Pt(10)
    contact_header_run.font.name = 'Calibri'
    contact_header_run.font.color.rgb = RGBColor(255, 255, 255)
    # contact_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # contact_header.paragraph_format.space_before = Pt(4)
    # contact_header.paragraph_format.space_after = Pt(4)
    # Use absolute path for icons
    icon_path = '/Users/prabhat/Desktop/resume'
    
    add_sidebar_text_with_icon(left_cell, os.path.join(icon_path, 'location.png'), 'Gurugram, Haryana', Pt(8.5))
    add_sidebar_text_with_icon(left_cell, os.path.join(icon_path, 'phone.png'), '9939147728', Pt(8.5))
    add_sidebar_text_with_icon(left_cell, os.path.join(icon_path, 'email.png'), 'prabhatkumar944@gmail.com', Pt(8))
    add_sidebar_text_with_icon(left_cell, os.path.join(icon_path, 'linkedin.png'), 'LinkedIn', Pt(8.5))
    add_sidebar_text_with_icon(left_cell, os.path.join(icon_path, 'github.png'), 'GitHub', Pt(8.5))
    add_sidebar_text_with_icon(left_cell, os.path.join(icon_path, 'portfolio.png'), 'Portfolio', Pt(8.5))
    
    # Technical Skills
    add_sidebar_header(left_cell, 'Technical Skills', is_main=False)
    add_sidebar_text(left_cell, 'Mobile Development', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'React Native, Expo, iOS, Android', Pt(8.5))
    add_sidebar_text(left_cell, 'Desktop Development', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'Electron, Tauri', Pt(8.5))
    add_sidebar_text(left_cell, 'Frameworks', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'React, Next.js, Node.js, Express.js, NestJS, Redux, Zustand', Pt(8.5))
    add_sidebar_text(left_cell, 'Languages', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'JavaScript, TypeScript, Kotlin, Swift, C#, Python', Pt(8.5))
    add_sidebar_text(left_cell, 'Databases', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'MongoDB, Redis, SQL, Elasticsearch', Pt(8.5))
    add_sidebar_text(left_cell, 'Tools & Cloud', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'Firebase, AWS, Docker, Git, Xcode, Android Studio, Visual Studio', Pt(8.5))
    
    # Education
    add_sidebar_header(left_cell, 'Education', is_main=False)
    add_sidebar_text(left_cell, 'Bachelor of Engineering', Pt(9), RGBColor(200, 220, 255))
    add_sidebar_text(left_cell, 'Electrical Engineering', Pt(8.5))
    add_sidebar_text(left_cell, 'Rajiv Gandhi Proudyogiki Vishwavidyalaya', Pt(8.5))
    add_sidebar_text(left_cell, '2015 - 2019', Pt(8.5))
    add_sidebar_text(left_cell, 'Bhopal, Madhya Pradesh', Pt(8.5))
    
    # Certifications
    add_sidebar_header(left_cell, 'Certifications', is_main=False)
    add_sidebar_bullet(left_cell, 'Node.js (Udemy)', Pt(8.5))
    add_sidebar_bullet(left_cell, 'MongoDB', Pt(8.5))
    add_sidebar_bullet(left_cell, 'React (HackerRank)', Pt(8.5))
    
    # Key Achievements
    add_sidebar_header(left_cell, 'Key Achievements', is_main=False)
    add_sidebar_bullet(left_cell, 'Published 5+ mobile apps to App Store & Play Store', Pt(8.5))
    add_sidebar_bullet(left_cell, 'Developed 3+ cross-platform desktop applications', Pt(8.5))
    add_sidebar_bullet(left_cell, 'Reduced crash rate by 40% across platforms', Pt(8.5))
    add_sidebar_bullet(left_cell, 'Improved user retention by 30%', Pt(8.5))
    add_sidebar_bullet(left_cell, 'Optimized server load by 30%', Pt(8.5))
    
    # ========== RIGHT COLUMN (MAIN CONTENT) ==========
    # Clear the default paragraph
    right_para = right_cell.paragraphs[0]
    right_para.clear()
    
    # Professional Summary
    summary_header = right_cell.add_paragraph()
    summary_header_run = summary_header.add_run('PROFESSIONAL SUMMARY')
    summary_header_run.bold = True
    summary_header_run.font.size = Pt(10.5)
    summary_header_run.font.name = 'Calibri'
    summary_header_run.font.color.rgb = RGBColor(0, 51, 102)
    summary_header.paragraph_format.space_before = Pt(0)
    summary_header.paragraph_format.space_after = Pt(1)
    summary_header.paragraph_format.line_spacing = 1.0
    summary_text = 'Desktop and Mobile Application Developer with 3+ years of experience building cross-platform iOS, Android, Windows, and macOS applications using React Native, Electron, Tauri, Kotlin, and Swift, with expertise in native integrations, performance optimization, and production-ready deployments.'
    summary_para = right_cell.add_paragraph()
    summary_run = summary_para.add_run(summary_text)
    summary_run.font.size = Pt(9.5)
    summary_run.font.name = 'Calibri'
    summary_run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
    summary_para.paragraph_format.line_spacing = 1.15
    summary_para.paragraph_format.space_after = Pt(5)
    
    # Work Experience
    add_main_header(right_cell, 'Work Experience')
    
    # Zupee
    exp1_company = right_cell.add_paragraph()
    exp1_company_run = exp1_company.add_run('Zupee')
    exp1_company_run.bold = True
    exp1_company_run.font.size = Pt(10.5)
    exp1_company_run.font.name = 'Calibri'
    exp1_company_run.font.color.rgb = RGBColor(0, 51, 102)
    exp1_company.paragraph_format.space_after = Pt(0)
    exp1_company.paragraph_format.line_spacing = 1.0
    
    exp1_line = right_cell.add_paragraph()
    exp1_line_run = exp1_line.add_run('Software Engineer - Mobile & Desktop App Development  |  Aug 2025 - Present  |  Gurugram, Haryana')
    exp1_line_run.font.size = Pt(9)
    exp1_line_run.font.name = 'Calibri'
    exp1_line_run.font.italic = True
    exp1_line_run.font.color.rgb = RGBColor(100, 100, 100)
    exp1_line.paragraph_format.space_before = Pt(0)
    exp1_line.paragraph_format.space_after = Pt(2)
    exp1_line.paragraph_format.line_spacing = 1.0
    
    add_compact_bullet(right_cell, 'Architected and developed 4 production-grade mobile applications using **React Native** with TypeScript, serving millions of users across iOS and Android platforms')
    add_compact_bullet(right_cell, 'Built cross-platform desktop applications using **Electron** and **Tauri**, ensuring consistent user experience across Windows, macOS, and Linux platforms')
    add_compact_bullet(right_cell, 'Built custom native bridge modules in **Kotlin** and **Swift** for FCM push notifications, universal deep linking, and integrated **HyperSDK** for **UPI** and other payment methods.')
    add_compact_bullet(right_cell, 'Implemented performance optimization strategies including lazy loading, code splitting, and memory leak detection, resulting in 50% faster app launch times across different platforms')
    
    # InsuranceDekho
    exp2_company = right_cell.add_paragraph()
    exp2_company_run = exp2_company.add_run('InsuranceDekho')
    exp2_company_run.bold = True
    exp2_company_run.font.size = Pt(10.5)
    exp2_company_run.font.name = 'Calibri'
    exp2_company_run.font.color.rgb = RGBColor(0, 51, 102)
    exp2_company.paragraph_format.space_before = Pt(4)
    exp2_company.paragraph_format.space_after = Pt(0)
    exp2_company.paragraph_format.line_spacing = 1.0
    
    exp2_line = right_cell.add_paragraph()
    exp2_line_run = exp2_line.add_run('Software Engineer  |  Mar 2025 - Jul 2025  |  Gurugram, Haryana')
    exp2_line_run.font.size = Pt(9)
    exp2_line_run.font.name = 'Calibri'
    exp2_line_run.font.italic = True
    exp2_line_run.font.color.rgb = RGBColor(100, 100, 100)
    exp2_line.paragraph_format.space_before = Pt(0)
    exp2_line.paragraph_format.space_after = Pt(2)
    exp2_line.paragraph_format.line_spacing = 1.0
    
    add_compact_bullet(right_cell, 'Built enterprise-grade web application using **Next.js** with Server-Side Rendering (SSR) and Static Site Generation (SSG) for optimal performance and SEO')
    add_compact_bullet(right_cell, 'Designed microservices architecture using **NestJS** with dependency injection, implementing RESTful APIs and GraphQL endpoints for flexible data access')
    add_compact_bullet(right_cell, 'Implemented distributed caching strategies with **Redis Cluster** and database query optimization, achieving 30% reduction in server load')
    add_compact_bullet(right_cell, 'Architected scalable data models in **MongoDB** with proper indexing strategies, aggregation pipelines, and transaction management for complex business logic')
    
    # Dresma
    exp3_company = right_cell.add_paragraph()
    exp3_company_run = exp3_company.add_run('Dresma')
    exp3_company_run.bold = True
    exp3_company_run.font.size = Pt(10.5)
    exp3_company_run.font.name = 'Calibri'
    exp3_company_run.font.color.rgb = RGBColor(0, 51, 102)
    exp3_company.paragraph_format.space_before = Pt(4)
    exp3_company.paragraph_format.space_after = Pt(0)
    exp3_company.paragraph_format.line_spacing = 1.0
    
    exp3_line = right_cell.add_paragraph()
    exp3_line_run = exp3_line.add_run('Software Engineer - Mobile App Development  |  Sep 2022 - Feb 2025  |  Gurugram, Haryana')
    exp3_line_run.font.size = Pt(9)
    exp3_line_run.font.name = 'Calibri'
    exp3_line_run.font.italic = True
    exp3_line_run.font.color.rgb = RGBColor(100, 100, 100)
    exp3_line.paragraph_format.space_before = Pt(0)
    exp3_line.paragraph_format.space_after = Pt(2)
    exp3_line.paragraph_format.line_spacing = 1.0
    
    add_compact_bullet(right_cell, 'Engineered cross-platform mobile application using **React Native** with native camera APIs, implementing real-time image processing algorithms and cloud storage integration')
    add_compact_bullet(right_cell, 'Designed scalable state management architecture using **Redux Toolkit** with **TanStack Query** for efficient API caching and normalized state structure')
    add_compact_bullet(right_cell, 'Integrated **Strapi** headless CMS for content management and **Segment** analytics platform for comprehensive user behavior tracking and data pipeline orchestration')
    add_compact_bullet(right_cell, 'Optimized bundle size through tree shaking, dynamic imports, and Hermes engine configuration, achieving 30% faster app startup and improved user retention')
    
    # Projects
    projects_header = right_cell.add_paragraph()
    projects_header_run = projects_header.add_run('COMPANY PROJECTS')
    projects_header_run.bold = True
    projects_header_run.font.size = Pt(10.5)
    projects_header_run.font.name = 'Calibri'
    projects_header_run.font.color.rgb = RGBColor(0, 51, 102)
    projects_header.paragraph_format.space_before = Pt(5)
    projects_header.paragraph_format.space_after = Pt(3)
    projects_header.paragraph_format.line_spacing = 1.0
    
    # Project 1 - Zupee Applications
    proj1_title = right_cell.add_paragraph()
    proj1_title_run = proj1_title.add_run('Mobile Applications: Ludo, Micro Drama Shots, Rummy, Astrology (iOS & Android)')
    proj1_title_run.bold = True
    proj1_title_run.font.size = Pt(9)
    proj1_title_run.font.name = 'Calibri'
    proj1_title_run.font.color.rgb = RGBColor(0, 51, 102)
    proj1_title.paragraph_format.space_after = Pt(0)
    proj1_title.paragraph_format.line_spacing = 1.0
    
    proj1_tech = right_cell.add_paragraph()
    proj1_tech_run = proj1_tech.add_run('React Native, Kotlin, Swift, Firebase, HyperSDK, Payment Gateways')
    proj1_tech_run.font.size = Pt(7.5)
    proj1_tech_run.font.name = 'Calibri'
    proj1_tech_run.font.italic = True
    proj1_tech_run.font.color.rgb = RGBColor(100, 100, 100)
    proj1_tech.paragraph_format.space_after = Pt(0.3)
    proj1_tech.paragraph_format.space_before = Pt(0)
    proj1_tech.paragraph_format.line_spacing = 1.0
    
    add_compact_bullet(right_cell, 'Architected modular mobile ecosystem with shared component libraries enabling code reusability across 4 applications')
    add_compact_bullet(right_cell, 'Implemented custom native bridge modules for FCM, universal links, and integrated **HyperSDK** for secure payment processing with multiple payment methods')
    
    # Project 2 - Image Editor Application
    proj2_title = right_cell.add_paragraph()
    proj2_title_run = proj2_title.add_run('Image Editor Application (iOS, Android, Desktop & Web)')
    proj2_title_run.bold = True
    proj2_title_run.font.size = Pt(9)
    proj2_title_run.font.name = 'Calibri'
    proj2_title_run.font.color.rgb = RGBColor(0, 51, 102)
    proj2_title.paragraph_format.space_before = Pt(1)
    proj2_title.paragraph_format.space_after = Pt(0)
    proj2_title.paragraph_format.line_spacing = 1.0
    
    proj2_tech = right_cell.add_paragraph()
    proj2_tech_run = proj2_tech.add_run('React Native, Electron, Tauri, React, Next.js, Redux, Redux Toolkit, Jotai, RxJS, Camera Integration, Image Processing')
    proj2_tech_run.font.size = Pt(7.5)
    proj2_tech_run.font.name = 'Calibri'
    proj2_tech_run.font.italic = True
    proj2_tech_run.font.color.rgb = RGBColor(100, 100, 100)
    proj2_tech.paragraph_format.space_after = Pt(0.3)
    proj2_tech.paragraph_format.space_before = Pt(0)
    proj2_tech.paragraph_format.line_spacing = 1.0
    
    add_compact_bullet(right_cell, 'Designed real-time image processing pipeline with GPU acceleration and custom rendering engine for cross-platform workflows across mobile, desktop, and web')
    add_compact_bullet(right_cell, 'Implemented state synchronization using **Redux Middleware**, **Jotai**, and **RxJS** for reactive data flow across all platforms')
    
    # Project 3 - SaaS Product
    proj3_title = right_cell.add_paragraph()
    proj3_title_run = proj3_title.add_run('SaaS Product (Web Application)')
    proj3_title_run.bold = True
    proj3_title_run.font.size = Pt(9)
    proj3_title_run.font.name = 'Calibri'
    proj3_title_run.font.color.rgb = RGBColor(0, 51, 102)
    proj3_title.paragraph_format.space_before = Pt(1)
    proj3_title.paragraph_format.space_after = Pt(0)
    proj3_title.paragraph_format.line_spacing = 1.0
    
    proj3_tech = right_cell.add_paragraph()
    proj3_tech_run = proj3_tech.add_run('React, Next.js, Node.js, NestJS, Redis, MongoDB, SaaS Architecture')
    proj3_tech_run.font.size = Pt(7.5)
    proj3_tech_run.font.name = 'Calibri'
    proj3_tech_run.font.italic = True
    proj3_tech_run.font.color.rgb = RGBColor(100, 100, 100)
    proj3_tech.paragraph_format.space_after = Pt(0.3)
    proj3_tech.paragraph_format.space_before = Pt(0)
    proj3_tech.paragraph_format.line_spacing = 1.0
    
    add_compact_bullet(right_cell, 'Designed multi-tenant SaaS architecture with tenant isolation, RBAC, and dynamic configuration management')
    add_compact_bullet(right_cell, 'Implemented event-driven microservices using **Message Queue** patterns, **Redis Pub/Sub**, and **MongoDB Change Streams**')
    
    # Save document
    output_file = 'Prabhat_Kumar_Resume_AppDeveloper.docx'
    doc.save(output_file)
    print(f"Two-column resume created successfully: {output_file}")

if __name__ == '__main__':
    try:
        create_twocolumn_resume()
    except ImportError:
        print("Error: python-docx library not found.")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error creating resume: {e}")
        import traceback
        traceback.print_exc()
