#!/usr/bin/env python3
"""
Script to create a two-column DOCX resume matching the Rajnish-style template:
centered header (name, title, contact with icons) and two white columns with
serif body text and sans-serif section headers. Design has no outer page border;
if you see a frame when exporting, disable border/frame in the export dialog.
Optional: for a very light grey/off-white background, set page color in Word
after opening the DOCX (Design > Page Color).

Output: Prabhat-Kumar-React.docx
To get Prabhat-Kumar-React.pdf, open Prabhat-Kumar-React.docx and export to PDF, or run
convert_to_pdf.py after editing it to use input Prabhat-Kumar-React.docx and output
Prabhat-Kumar-React.pdf.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
import os

# Script-relative base path for icons
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Fonts: Use Times New Roman throughout (serif font matching the design)
SERIF_FONT = 'Times New Roman'
SANS_SERIF_FONT = 'Times New Roman'  # Use same font throughout

# Dark teal-green for section header icons and borders (matches reference image)
SECTION_ACCENT_COLOR = '2D5A5A'  # OOXML hex (no #)
SECTION_ACCENT_COLOR_URL = '%232D5A5A'  # URL-encoded for Iconify

# Icon file mappings - PNG files in icon folder
ICON_FILES = {
    'profile': 'profile.png',
    'skills': 'skill.png',
    'education': 'education.png',
    'certificates': 'certificate.png',
    'findmeonline': 'findme.png',
    'experience': 'experience.png',
    'projects': 'project.png',
    'external_link': 'link.png',  # Using link.png as external link icon
    'location': 'location-pin.png',
    'email': 'email.png',
    'phone': 'call.png',
}

def get_icon_path(icon_key):
    """Get full path to icon PNG file."""
    if icon_key not in ICON_FILES:
        return None
    icon_file = ICON_FILES[icon_key]
    icon_path = os.path.join(SCRIPT_DIR, 'icon', icon_file)
    if os.path.exists(icon_path):
        return icon_path
    return None

def set_margins(section, top, bottom, left, right):
    """Set page margins for a section."""
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def _add_short_bold_line(cell):
    """Add a short, bold horizontal line below section header."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)  # No gap - line directly below title
    # p.paragraph_format.space_after = Pt(4)
    # p.paragraph_format.line_spacing = 1.0  # Tight line spacing
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Inches(3)  # Make it shorter (not full width)
    
    # Add bold bottom border
    pPr = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Thicker line (12 = 1.5pt, was 6 = 0.75pt)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SECTION_ACCENT_COLOR)  # Dark teal-green to match icons
    pBdr.append(bottom)
    pPr.append(pBdr)
    p._element.insert(0, pPr)
    return p

def add_section_header(cell, text, icon_key=None, is_left_column=False):
    """Add a section header: optional icon, bold uppercase sans-serif title, then short bold line."""
    # Text paragraph - no right indent to prevent vertical wrapping
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)  # No gap before border line
    p.paragraph_format.line_spacing = 1.0  # Tight line spacing
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    # NO right_indent here - it causes text to wrap vertically
    
    if icon_key:
        icon_path = get_icon_path(icon_key)
        if icon_path:
            try:
                # Ensure no left indent for both columns
                pPr = p._element.get_or_add_pPr()
                # Remove existing indentation if present
                existing_ind = pPr.find(qn('w:ind'))
                if existing_ind is not None:
                    pPr.remove(existing_ind)
                
                # Only apply negative indent for left column to offset its cell margin
                # Right column has 0 left margin, so no negative indent needed
                if is_left_column:
                    # Create indentation element with smaller negative indent to move icons more to the right
                    # -144 twips = -0.1 inches (smaller offset to move icons more to the right)
                    # This keeps icons away from the left edge while maintaining proper alignment
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), '-144')  # Smaller negative indent to move icons more to the right
                    ind.set(qn('w:firstLine'), '0')
                    pPr.append(ind)
                else:
                    # For right column, explicitly set left indent to 0 to remove any default spacing
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), '-144')
                    ind.set(qn('w:firstLine'), '0')
                    pPr.append(ind)
                
                icon_run = p.add_run()
                icon_run.add_picture(icon_path, width=Pt(10), height=Pt(10))
                # Minimal spacing between icon and text - using very small non-breaking space
                sp = p.add_run('\u2009')  # Thin space character (smaller than regular space)
                sp.font.size = Pt(0.5)  # Very small spacing to minimize gap
            except Exception as e:
                print(f"Warning: Could not add icon {icon_key}: {e}")
                pass
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = SANS_SERIF_FONT  # Use same font as all other content
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Border line paragraph - separate paragraph with right indent for shorter line
    border_p = cell.add_paragraph()
    border_p.paragraph_format.space_before = Pt(0)  # No gap - directly below title
    border_p.paragraph_format.space_after = Pt(2)  # Space after the border line (compact for one page)
    border_p.paragraph_format.line_spacing = 1.0
    border_p.paragraph_format.left_indent = Pt(0)
    # Adjust right indent based on column - smaller right_indent = longer border line
    if is_left_column:
        border_p.paragraph_format.right_indent = Inches(0.5)  # Left column: longer border below title
    else:
        border_p.paragraph_format.right_indent = Inches(0.5)  # Right column: border to match icon+title
    
    # Add a non-breaking space to give the paragraph width so the border renders properly
    border_run = border_p.add_run('\u00A0')  # Non-breaking space
    border_run.font.size = Pt(1)  # Very small, invisible content
    
    # Add bold bottom border to this paragraph
    pPr = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Thicker line (12 = 1.5pt)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SECTION_ACCENT_COLOR)  # Dark teal-green to match icons
    pBdr.append(bottom)
    pPr.append(pBdr)
    border_p._element.insert(0, pPr)
    
    return p

def add_body_text(cell, text, size=Pt(9)):
    """Add body paragraph, black text (serif)."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = SANS_SERIF_FONT
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    return p

def add_bullet(cell, text, size=Pt(9)):
    """Add bullet point with optional **bold** keywords (bold in black)."""
    p = cell.add_paragraph()
    bullet_run = p.add_run('• ')
    bullet_run.font.size = size
    bullet_run.font.name = SANS_SERIF_FONT
    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
    bullet_run.bold = True
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.line_spacing = 1.0
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.size = size
                run.font.name = SANS_SERIF_FONT
                run.font.color.rgb = RGBColor(0, 0, 0)
                if i % 2 == 1:
                    run.bold = True
    else:
        run = p.add_run(text)
        run.font.size = size
        run.font.name = SANS_SERIF_FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_contact_header_centered(cell_or_doc, icon_dir, location, email, phone, icon_size=Pt(9)):
    """Add centered contact line with inline icons (location, email, phone)."""
    p = cell_or_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    items = [
        ('location', location),
        ('email', email),
        ('phone', phone),
    ]
    for i, (icon_key, text) in enumerate(items):
        icon_path = get_icon_path(icon_key)
        if icon_path:
            try:
                icon_run = p.add_run()
                icon_run.add_picture(icon_path, width=icon_size, height=icon_size)
                # Reduced spacing between icon and text
                sp = p.add_run('\u2009')  # Thin space character
                sp.font.size = Pt(2)  # Reduced from Pt(6) to minimize gap
            except Exception as e:
                print(f"Warning: Could not add icon {icon_key}: {e}")
                pass
        txt = p.add_run(text)
        txt.font.size = Pt(10)
        txt.font.name = SANS_SERIF_FONT
        txt.font.color.rgb = RGBColor(0, 0, 0)
        if '@' in text:
            txt.element.set(qn('w:noBreak'), '1')
        if i < len(items) - 1:
            sep = p.add_run('  ')  # Reduced spacing between contact items from '   ' to '  '
            sep.font.size = Pt(4)  # Reduced from Pt(6)
    return p

def set_cell_margins(cell, top, bottom, left, right):
    """Set cell margins (inches)."""
    tc = cell._element
    tcPr = tc.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.append(tcPr)
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(int(margin_value * 1440)))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def create_prabhat_resume():
    """Build the Rajnish-style two-column resume for Prabhat Kumar as React Developer (no page border; light grey page background)."""
    doc = Document()
    section = doc.sections[0]
    set_margins(section, 0, 0, 0, 0)  # No page margins - content fits end-to-end
    
    # Set page background to very light grey
    sectPr = section._sectPr
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), 'F8F8F8')  # Very light grey
    sectPr.append(bg)

    style = doc.styles['Normal']
    style.font.name = SANS_SERIF_FONT
    style.font.size = Pt(10.5)

    # ----- Header: name, title, contact (with background color) -----
    # Create header table with background color
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Set table width to full page width (8.5" page)
    header_table.columns[0].width = Inches(8.5)

    # Set background color for header cell
    header_tcPr = header_cell._element.tcPr
    if header_tcPr is None:
        header_tcPr = OxmlElement('w:tcPr')
        header_cell._element.insert(0, header_tcPr)
    header_shading = OxmlElement('w:shd')
    header_shading.set(qn('w:fill'), 'F5F5F5')  # Very light grey
    header_tcPr.append(header_shading)

    # Remove table borders
    header_tbl = header_table._element
    header_tblPr = header_tbl.tblPr
    if header_tblPr is None:
        header_tblPr = OxmlElement('w:tblPr')
        header_tbl.insert(0, header_tblPr)
    header_tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        header_tblBorders.append(border)
    header_tblPr.append(header_tblBorders)

    # Set cell margins - compact for one-page resume
    set_cell_margins(header_cell, 0.04, 0.12, 0.0, 0.0)

    # Clear default paragraph and add header content
    header_cell.paragraphs[0].clear()

    name_p = header_cell.add_paragraph()
    name_p.paragraph_format.space_before = Pt(4)  # Add some top spacing
    name_run = name_p.add_run('Prabhat Kumar')
    name_run.bold = True
    name_run.font.size = Pt(19)
    name_run.font.name = SANS_SERIF_FONT
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)

    title_p = header_cell.add_paragraph()
    title_run = title_p.add_run('Software Engineer')
    title_run.font.size = Pt(12)
    title_run.font.name = SANS_SERIF_FONT
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(3)

    add_contact_header_centered(
        header_cell, SCRIPT_DIR,
        'Gurugram, Haryana',
        'prabhatkumar944@gmail.com',
        '+91 9939147728',
    )
    
    # Add spacing after header (compact for one page)
    if header_cell.paragraphs:
        header_cell.paragraphs[-1].paragraph_format.space_after = Pt(4)

    # ----- Two-column table -----
    table = doc.add_table(rows=1, cols=2)
    left_col = table.columns[0]
    right_col = table.columns[1]
    # Full page width (8.5") - equal 50/50 split
    left_col.width = Inches(2.40)
    right_col.width = Inches(6.10)

    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    # Add left padding to prevent text cutoff, minimal other margins
    set_cell_margins(left_cell, 0.0, 0.0, 0.45, 0.2)  # Increased left padding
    set_cell_margins(right_cell, 0.0, 0.0, 0.0, 0.1)  # Small left padding, no right margin
    
    # Set white background for both columns to contrast with page background
    left_tcPr = left_cell._element.tcPr
    if left_tcPr is None:
        left_tcPr = OxmlElement('w:tcPr')
        left_cell._element.insert(0, left_tcPr)
    shading_left = OxmlElement('w:shd')
    shading_left.set(qn('w:fill'), 'FFFFFF')  # White
    left_tcPr.append(shading_left)
    
    right_tcPr = right_cell._element.tcPr
    if right_tcPr is None:
        right_tcPr = OxmlElement('w:tcPr')
        right_cell._element.insert(0, right_tcPr)
    shading_right = OxmlElement('w:shd')
    shading_right.set(qn('w:fill'), 'FFFFFF')  # White
    right_tcPr.append(shading_right)

    tr = table.rows[0]._element
    trPr = tr.trPr
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '0')
    trHeight.set(qn('w:hRule'), 'auto')
    trPr.append(trHeight)

    # Clear default paragraph in cells
    left_cell.paragraphs[0].clear()
    right_cell.paragraphs[0].clear()

    # ========== LEFT COLUMN ==========
    add_section_header(left_cell, 'Profile', icon_key='profile', is_left_column=True)
    add_body_text(
        left_cell,
        'React Developer with 3+ years building cross-platform mobile and web apps using React Native, React, and TypeScript. Strong in state management, component architecture, and responsive frontend development. Experienced in agile delivery, code reviews, and building reusable component libraries for scalable applications.',
        size=Pt(11)
    )

    add_section_header(left_cell, 'Skills', icon_key='skills', is_left_column=True)
    # Frameworks & Libraries
    row1 = left_cell.add_paragraph()
    row1_prefix = row1.add_run('Frameworks & Libraries:')
    row1_prefix.bold = True
    row1_prefix.font.size = Pt(10.5)
    row1_prefix.font.name = SANS_SERIF_FONT
    row1_prefix.font.color.rgb = RGBColor(0, 0, 0)
    row1_content = row1.add_run(' React, Next.js, React Native, Redux, Redux Toolkit, Zustand, Expo')
    row1_content.font.size = Pt(10.5)
    row1_content.font.name = SANS_SERIF_FONT
    row1_content.font.color.rgb = RGBColor(0, 0, 0)
    row1.paragraph_format.space_after = Pt(1)

    # Languages
    row2 = left_cell.add_paragraph()
    row2_prefix = row2.add_run('Languages:')
    row2_prefix.bold = True
    row2_prefix.font.size = Pt(10.5)
    row2_prefix.font.name = SANS_SERIF_FONT
    row2_prefix.font.color.rgb = RGBColor(0, 0, 0)
    row2_content = row2.add_run(' JavaScript, TypeScript, HTML, CSS')
    row2_content.font.size = Pt(10.5)
    row2_content.font.name = SANS_SERIF_FONT
    row2_content.font.color.rgb = RGBColor(0, 0, 0)
    row2.paragraph_format.space_after = Pt(1)

    # Features & Integration
    row3 = left_cell.add_paragraph()
    row3_prefix = row3.add_run('Features & Integration:')
    row3_prefix.bold = True
    row3_prefix.font.size = Pt(10.5)
    row3_prefix.font.name = SANS_SERIF_FONT
    row3_prefix.font.color.rgb = RGBColor(0, 0, 0)
    row3_content = row3.add_run(' Payment Integration, Push Notifications, Deep Linking, Native Navigation')
    row3_content.font.size = Pt(10.5)
    row3_content.font.name = SANS_SERIF_FONT
    row3_content.font.color.rgb = RGBColor(0, 0, 0)
    row3.paragraph_format.space_after = Pt(1)

    # Tools
    row4 = left_cell.add_paragraph()
    row4_prefix = row4.add_run('Tools:')
    row4_prefix.bold = True
    row4_prefix.font.size = Pt(10.5)
    row4_prefix.font.name = SANS_SERIF_FONT
    row4_prefix.font.color.rgb = RGBColor(0, 0, 0)
    row4_content = row4.add_run(' Git, Firebase, Figma, JIRA, Postman, Docker, Android Studio, Xcode')
    row4_content.font.size = Pt(10.5)
    row4_content.font.name = SANS_SERIF_FONT
    row4_content.font.color.rgb = RGBColor(0, 0, 0)
    row4.paragraph_format.space_after = Pt(1)

    add_section_header(left_cell, 'Education', icon_key='education', is_left_column=True)
    # B.Tech
    btech_label = left_cell.add_paragraph()
    btech_label_run = btech_label.add_run('Bachelor of Engineering')
    btech_label_run.bold = True
    btech_label_run.font.size = Pt(10.5)
    btech_label_run.font.name = SANS_SERIF_FONT
    btech_label_run.font.color.rgb = RGBColor(0, 0, 0)
    btech_label.paragraph_format.space_after = Pt(0.5)
    add_body_text(left_cell, 'Rajiv Gandhi Proudyogiki Vishwavidyalaya | 2015 – 2019', Pt(10.5))
    add_body_text(left_cell, 'B.E. Electrical Engineering, Bhopal', Pt(10.5))

    add_section_header(left_cell, 'Certificates', icon_key='certificates', is_left_column=True)
    add_body_text(left_cell, 'React (HackerRank)', Pt(10.5))


    # ========== RIGHT COLUMN ==========
    add_section_header(right_cell, 'Professional Experience', icon_key='experience', is_left_column=False)

    # Zupee - Software Developer (Most recent)
    exp1_company = right_cell.add_paragraph()
    exp1_run = exp1_company.add_run('Zupee')
    exp1_run.bold = True
    exp1_run.font.size = Pt(11)
    exp1_run.font.name = SANS_SERIF_FONT
    exp1_run.font.color.rgb = RGBColor(0, 0, 0)
    exp1_company.paragraph_format.space_after = Pt(0)
    
    exp1_title = right_cell.add_paragraph()
    exp1_title_run = exp1_title.add_run('Software Developer')
    exp1_title_run.bold = True
    exp1_title_run.font.size = Pt(10)
    exp1_title_run.font.name = SANS_SERIF_FONT
    exp1_title_run.font.color.rgb = RGBColor(0, 0, 0)
    exp1_title.paragraph_format.space_after = Pt(0)
    
    exp1_line = right_cell.add_paragraph()
    exp1_line_run = exp1_line.add_run('Aug 2025 – Present | Gurugram, Haryana')
    exp1_line_run.font.size = Pt(10)
    exp1_line_run.font.name = SANS_SERIF_FONT
    exp1_line_run.font.color.rgb = RGBColor(80, 80, 80)
    exp1_line.paragraph_format.space_after = Pt(1)
    add_bullet(right_cell, 'Architected and developed 4 production-grade mobile applications using **React Native** with TypeScript, serving millions of users across iOS and Android platforms', size=Pt(11))
    add_bullet(right_cell, 'Built scalable **React** web applications with **Redux** and **Zustand** for state management, ensuring optimal performance and user experience', size=Pt(11))
    add_bullet(right_cell, 'Implemented custom native bridge modules in **Kotlin** and **Swift** for FCM push notifications, universal deep linking, and payment gateway integrations', size=Pt(11))
    add_bullet(right_cell, 'Optimized React Native applications through code splitting, lazy loading, and memory management, achieving 50% faster app launch times', size=Pt(11))

    # InsuranceDekho - Software Developer
    exp2_company = right_cell.add_paragraph()
    exp2_run = exp2_company.add_run('InsuranceDekho')
    exp2_run.bold = True
    exp2_run.font.size = Pt(11)
    exp2_run.font.name = SANS_SERIF_FONT
    exp2_run.font.color.rgb = RGBColor(0, 0, 0)
    exp2_company.paragraph_format.space_before = Pt(2)
    exp2_company.paragraph_format.space_after = Pt(0)
    
    exp2_title = right_cell.add_paragraph()
    exp2_title_run = exp2_title.add_run('Software Developer')
    exp2_title_run.bold = True
    exp2_title_run.font.size = Pt(10)
    exp2_title_run.font.name = SANS_SERIF_FONT
    exp2_title_run.font.color.rgb = RGBColor(0, 0, 0)
    exp2_title.paragraph_format.space_after = Pt(0)
    
    exp2_line = right_cell.add_paragraph()
    exp2_line_run = exp2_line.add_run('Mar 2025 – Jul 2025 | Gurugram, Haryana')
    exp2_line_run.font.size = Pt(10)
    exp2_line_run.font.name = SANS_SERIF_FONT
    exp2_line_run.font.color.rgb = RGBColor(80, 80, 80)
    exp2_line.paragraph_format.space_after = Pt(1)
    add_bullet(right_cell, 'Developed and deployed **insurance booking portal** screens; integrated multiple **third-party** kits to fetch quotes and enabled users to **compare** policies using filters (2W, 4W, **health**, **motor**, and other policy types) and complete booking', size=Pt(11))
    add_bullet(right_cell, 'Built enterprise-grade web application using **Next.js** with SSR and SSG for performance and SEO; designed reusable **React** component architecture with TypeScript', size=Pt(11))
    add_bullet(right_cell, 'Implemented state management with **Redux** and **Zustand**; integrated REST and **GraphQL** APIs for data fetching and error handling', size=Pt(11))
    add_bullet(right_cell, 'Delivered **responsive** UI with **TypeScript**; ensured **accessibility** and consistent UX across quote comparison and policy booking flows', size=Pt(11))

    # Dresma - Software Developer
    exp3_company = right_cell.add_paragraph()
    exp3_run = exp3_company.add_run('Dresma')
    exp3_run.bold = True
    exp3_run.font.size = Pt(11)
    exp3_run.font.name = SANS_SERIF_FONT
    exp3_run.font.color.rgb = RGBColor(0, 0, 0)
    exp3_company.paragraph_format.space_before = Pt(2)
    exp3_company.paragraph_format.space_after = Pt(0)
    
    exp3_title = right_cell.add_paragraph()
    exp3_title_run = exp3_title.add_run('Software Developer')
    exp3_title_run.bold = True
    exp3_title_run.font.size = Pt(10)
    exp3_title_run.font.name = SANS_SERIF_FONT
    exp3_title_run.font.color.rgb = RGBColor(0, 0, 0)
    exp3_title.paragraph_format.space_after = Pt(0)
    
    exp3_line = right_cell.add_paragraph()
    exp3_line_run = exp3_line.add_run('Sep 2022 – Feb 2025 | Gurugram, Haryana')
    exp3_line_run.font.size = Pt(10)
    exp3_line_run.font.name = SANS_SERIF_FONT
    exp3_line_run.font.color.rgb = RGBColor(80, 80, 80)
    exp3_line.paragraph_format.space_after = Pt(1)
    add_bullet(right_cell, 'Built web-based editor using **React**, **Next.js**, **Fabric.js**, and **RxJS** for canvas manipulation and reactive data flow', size=Pt(11))
    add_bullet(right_cell, 'Designed central state architecture with **Zustand**, **Redux**, and **Jotai** for scalable store management and component synchronization', size=Pt(11))
    add_bullet(right_cell, 'Implemented API layer using **TanStack Query** (useQuery) for server state, caching, and **batch API** calls; optimized request batching and invalidation', size=Pt(11))
    add_bullet(right_cell, 'Applied **lazy loading**, code splitting, and performance optimizations to improve load times and runtime efficiency of the editor application', size=Pt(11))

    add_section_header(right_cell, 'Projects', icon_key='projects', is_left_column=False)

    # Project 1 - Image Editor (Web)
    proj1_title = right_cell.add_paragraph()
    proj1_run = proj1_title.add_run('Image Editor (Web)')
    proj1_run.bold = True
    proj1_run.font.size = Pt(11)
    proj1_run.font.name = SANS_SERIF_FONT
    proj1_run.font.color.rgb = RGBColor(0, 0, 0)
    proj1_title.paragraph_format.space_before = Pt(2)
    proj1_title.paragraph_format.space_after = Pt(0)
    
    proj1_tech = right_cell.add_paragraph()
    tech_label_run1 = proj1_tech.add_run('Tech Stack:')
    tech_label_run1.bold = True
    tech_label_run1.font.size = Pt(10)
    tech_label_run1.font.name = SANS_SERIF_FONT
    tech_label_run1.font.color.rgb = RGBColor(0, 0, 0)
    tech_content_run1 = proj1_tech.add_run(' React, Next.js, Fabric.js, RxJS, Zustand, Redux, Jotai, TanStack Query')
    tech_content_run1.font.size = Pt(10)
    tech_content_run1.font.name = SANS_SERIF_FONT
    tech_content_run1.font.color.rgb = RGBColor(0, 0, 0)
    proj1_tech.paragraph_format.space_after = Pt(1)
    
    add_bullet(right_cell, 'Web-based canvas editor with **Fabric.js** and **RxJS**; central state via **Zustand**, **Redux**, **Jotai**; **TanStack Query** for API caching; lazy loading and code splitting', size=Pt(11))

    # Project 2 - Enterprise Web Application
    proj2_title = right_cell.add_paragraph()
    proj2_run = proj2_title.add_run('Enterprise Web Application (Next.js, React)')
    proj2_run.bold = True
    proj2_run.font.size = Pt(11)
    proj2_run.font.name = SANS_SERIF_FONT
    proj2_run.font.color.rgb = RGBColor(0, 0, 0)
    proj2_title.paragraph_format.space_before = Pt(2)
    proj2_title.paragraph_format.space_after = Pt(0)
    
    proj2_tech = right_cell.add_paragraph()
    tech_label_run2 = proj2_tech.add_run('Tech Stack:')
    tech_label_run2.bold = True
    tech_label_run2.font.size = Pt(10)
    tech_label_run2.font.name = SANS_SERIF_FONT
    tech_label_run2.font.color.rgb = RGBColor(0, 0, 0)
    tech_content_run2 = proj2_tech.add_run(' React, Next.js, TypeScript, Redux, Zustand')
    tech_content_run2.font.size = Pt(10)
    tech_content_run2.font.name = SANS_SERIF_FONT
    tech_content_run2.font.color.rgb = RGBColor(0, 0, 0)
    proj2_tech.paragraph_format.space_after = Pt(1)
    
    add_bullet(right_cell, 'Built **Next.js** web app with SSR/SSG; reusable **React** components; **Redux**/Zustand state management', size=Pt(11))

    # Project 3 - SaaS Web Application
    proj3_title = right_cell.add_paragraph()
    proj3_run = proj3_title.add_run('SaaS Product (Web Application)')
    proj3_run.bold = True
    proj3_run.font.size = Pt(11)
    proj3_run.font.name = SANS_SERIF_FONT
    proj3_run.font.color.rgb = RGBColor(0, 0, 0)
    proj3_title.paragraph_format.space_before = Pt(2)
    proj3_title.paragraph_format.space_after = Pt(0)
    
    proj3_tech = right_cell.add_paragraph()
    tech_label_run3 = proj3_tech.add_run('Tech Stack:')
    tech_label_run3.bold = True
    tech_label_run3.font.size = Pt(10)
    tech_label_run3.font.name = SANS_SERIF_FONT
    tech_label_run3.font.color.rgb = RGBColor(0, 0, 0)
    tech_content_run3 = proj3_tech.add_run(' React, Next.js, Redux, TypeScript')
    tech_content_run3.font.size = Pt(10)
    tech_content_run3.font.name = SANS_SERIF_FONT
    tech_content_run3.font.color.rgb = RGBColor(0, 0, 0)
    proj3_tech.paragraph_format.space_after = Pt(1)
    
    add_bullet(right_cell, 'Multi-tenant SaaS UI with **React** frontend, **Next.js** SSR, **Redux** state management', size=Pt(11))

    # Save
    output_file = os.path.join(SCRIPT_DIR, 'Prabhat-Kumar-React.docx')
    doc.save(output_file)
    print(f"Resume created: {output_file}")
    print("To get Prabhat-Kumar-React.pdf, open the DOCX in Word and export to PDF, or run convert_to_pdf.py with this file.")


if __name__ == '__main__':
    try:
        create_prabhat_resume()
    except ImportError:
        print("Error: python-docx not found. Install with: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
